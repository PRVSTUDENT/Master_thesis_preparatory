from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
STAGE = ROOT / "runs" / "chaboche_umat" / "stage16_abaqus_inhomogeneous_cycle_jump_benchmark"
OUT_DIR = ROOT / "docs" / "reports" / "stage16n_geometry_workflow"
DOCX = OUT_DIR / "stage16n_geometry_and_cycle_jump_workflow.docx"
FLOW = OUT_DIR / "stage16n_cycle_jump_workflow_flowchart.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    return p


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.25)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Item", True)
    set_cell_text(hdr[1], "Stage 16N setting", True)
    set_cell_shading(hdr[0], "F2F4F7")
    set_cell_shading(hdr[1], "F2F4F7")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def make_flowchart():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    steps = [
        ("1. Freeze material model", "Use fixed NEML-equivalent Chaboche UMAT."),
        ("2. Build inhomogeneous benchmark", "Plate with central hole, C3D8 mesh, cyclic U1 loading."),
        ("3. Run full reference", "1000 non-jump Abaqus cycles completed."),
        ("4. Confirm evolution", "RF loop area and local STATEV evolution exceed thresholds."),
        ("5. Set efficient CPU policy", "Use 16 CPUs by default; verify x 16 THREADS."),
        ("6. Analyze reference data", "Extract RF peaks, loop area, stresses, and STATEVs."),
        ("7. Compute adaptive Delta N", "Use most restrictive monitored variable."),
        ("8. Validate fixed jumps", "Compare controlled jumps inside 1000-cycle reference."),
        ("9. Implement adaptive jumping", "Extrapolate and reinject state, then correct/re-evaluate."),
        ("10. Final comparison", "Accuracy versus cost saving against full reference."),
    ]

    fig, ax = plt.subplots(figsize=(8.5, 11), dpi=180)
    ax.axis("off")
    x = 0.5
    y_top = 0.94
    box_w = 0.78
    box_h = 0.075
    gap = 0.018
    for i, (title, detail) in enumerate(steps):
        y = y_top - i * (box_h + gap)
        color = "#E8EEF5" if i < 5 else "#F4F6F9"
        rect = plt.Rectangle(
            (x - box_w / 2, y - box_h / 2),
            box_w,
            box_h,
            facecolor=color,
            edgecolor="#2E74B5",
            linewidth=1.2,
        )
        ax.add_patch(rect)
        ax.text(x, y + 0.013, title, ha="center", va="center", fontsize=10.5, fontweight="bold", color="#0B2545")
        ax.text(x, y - 0.017, detail, ha="center", va="center", fontsize=8.8, color="#222222")
        if i < len(steps) - 1:
            ax.annotate(
                "",
                xy=(x, y - box_h / 2 - gap * 0.25),
                xytext=(x, y - box_h / 2 - gap * 0.95),
                arrowprops=dict(arrowstyle="<-", color="#2E74B5", lw=1.0),
            )
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout(pad=0.25)
    fig.savefig(FLOW)
    plt.close(fig)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_doc():
    make_flowchart()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Stage 16N Geometry and Cycle-Jump Workflow")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run("Plate-with-hole Abaqus benchmark using a NEML-equivalent Chaboche UMAT")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document summarizes the Stage 16N benchmark geometry and the planned workflow for fixed and adaptive "
        "cycle-jump validation. The material model is kept fixed; the thesis contribution is the cycle-jump methodology "
        "and its accuracy-versus-cost evaluation on an inhomogeneous Abaqus model."
    )

    doc.add_heading("Geometry Used in Stage 16N", level=1)
    add_key_value_table(
        doc,
        [
            ("Benchmark type", "Thin 3D plate with a central circular hole."),
            ("Dimensions", "Length = 20.0, height = 10.0, thickness = 1.0."),
            ("Hole", "Central hole with radius = 1.0 at the plate center."),
            ("Mesh spacing", "Structured grid with dx = 0.25 and dy = 0.25."),
            ("Element type", "C3D8 for Stage 16N NEML-equivalent Abaqus deck."),
            ("Nodes", "6642 generated nodes."),
            ("Elements", "3148 solid elements outside the hole."),
            ("Hole-ring set", "60 elements selected near the hole for local stress and STATEV tracking."),
            ("Boundary sets", "82 left-edge nodes and 82 right-edge nodes."),
            ("Anchors", "Anchor A = node 1; Anchor B = node 3241, used to remove rigid-body motion."),
        ],
    )

    geom = STAGE / "stage16_plate_with_hole_geometry.png"
    mesh = STAGE / "stage16_plate_with_hole_actual_mesh.png"
    if geom.exists():
        doc.add_picture(str(geom), width=Inches(5.8))
        add_caption(doc, "Figure 1. Stage 16 plate-with-hole geometry schematic.")
    if mesh.exists():
        doc.add_picture(str(mesh), width=Inches(5.8))
        add_caption(doc, "Figure 2. Generated plate-with-hole mesh used for the benchmark.")

    doc.add_heading("Loading, Material, and Output", level=1)
    add_key_value_table(
        doc,
        [
            ("Material model", "NEML-equivalent Chaboche cyclic plasticity UMAT."),
            ("State variables", "27 solution-dependent state variables in the Stage 16N UMAT."),
            ("Elastic constants", "E = 200000.0, nu = 0.3."),
            ("Chaboche-style parameters", "yield = 100.0, Q = 50.0, b = 5.0, three backstress pairs: C1/g1, C2/g2, C3/g3."),
            ("Cycle definition", "Separate static step for each cycle, with AMP_ONE_CYCLE tabular loading."),
            ("Right edge loading", "Cyclic displacement U1 amplitude = +/-0.10 on RIGHT_EDGE."),
            ("Left edge constraint", "LEFT_EDGE fixed in U1; anchor nodes constrain remaining rigid-body modes."),
            ("Reference run", "1000 full non-jump cycles completed as the accepted baseline."),
            ("Selected field cycles", "1, 2, 10, 50, 100, 250, 500, 750, and 1000."),
            ("Tracked outputs", "RF1-U1 hysteresis, loop area, local hole-ring stresses, and selected STATEV values."),
        ],
    )

    doc.add_heading("Reference Status", level=1)
    add_bullet(doc, "The corrected full reference completed 1000 non-jump cycles.")
    add_bullet(doc, "Abaqus parallelism was verified as 1 MPI rank x 30 threads for the completed full reference.")
    add_bullet(doc, "Future production runs are configured to use 16 CPUs by default to reduce resource waste.")
    add_bullet(doc, "The reference CSV files provide the baseline for fixed and adaptive cycle-jump validation.")

    doc.add_heading("Plan Flow Chart", level=1)
    doc.add_picture(str(FLOW), width=Inches(6.1))
    add_caption(doc, "Figure 3. Proposed Stage 16N workflow from benchmark setup to final cycle-jump validation.")

    doc.add_heading("Key Decision Points", level=1)
    add_key_value_table(
        doc,
        [
            ("Benchmark suitability", "Accepted: global loop evolution and local hole STATEV evolution are significant."),
            ("Early-cycle treatment", "Do not jump the first cycles aggressively; simulate the transient normally."),
            ("Delta N control", "Use local STATEV and local stress variables in addition to global RF1 and loop area."),
            ("Validation order", "Run fixed controlled jumps before the full adaptive cycle-jump workflow."),
            ("CPU policy", "Use 16 CPUs by default and verify the Abaqus message file reports x 16 THREADS."),
        ],
    )

    doc.add_heading("Immediate Next Tasks", level=1)
    for item in [
        "Finish/check the 16-CPU benchmark and record the measured efficiency.",
        "Analyze the 1000-cycle full reference data.",
        "Compute the adaptive Delta N table from the full reference.",
        "Run fixed jump validation cases inside the 1000-cycle window.",
        "Implement adaptive cycle jumping after the fixed jumps pass the accuracy checks.",
    ]:
        add_bullet(doc, item)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build_doc()
